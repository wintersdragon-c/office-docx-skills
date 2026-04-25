from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def m_tag(local_name: str) -> str:
    return f"{{{MATH_NS}}}{local_name}"


def make_text_formula_omml(formula_text: str):
    """Build a minimal editable Word equation containing plain math text."""
    o_math = OxmlElement("m:oMath")
    run = OxmlElement("m:r")
    text = OxmlElement("m:t")
    text.text = formula_text
    run.append(text)
    o_math.append(run)
    return o_math


def add_omml_formula(paragraph, formula_text: str) -> None:
    """Append a minimal OMML equation to a python-docx paragraph."""
    paragraph._p.append(make_text_formula_omml(formula_text))


def create_docx_with_omml_formula(
    output_path: str | Path,
    formula_text: str,
    body_text: str | None = None,
) -> None:
    """Create a DOCX containing one native Word equation for smoke tests/examples."""
    document = Document()
    if body_text:
        document.add_paragraph(body_text)
    formula_paragraph = document.add_paragraph()
    formula_paragraph.style = document.styles["Normal"]
    add_omml_formula(formula_paragraph, formula_text)

    # Keep the formula font explicit in the document defaults where Word supports it.
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")

    document.save(output_path)
