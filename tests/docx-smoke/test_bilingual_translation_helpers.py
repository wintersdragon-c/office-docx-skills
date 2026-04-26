#!/usr/bin/env python3
from __future__ import annotations

import importlib.util
import sys
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

ROOT = Path(__file__).resolve().parents[2]
HELPER = ROOT / "skills" / "docx-bilingual-translation" / "translation_docx_helpers.py"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
V = "urn:schemas-microsoft-com:vml"
A = "http://schemas.openxmlformats.org/drawingml/2006/main"
NS = {"w": W, "m": M, "r": R, "v": V, "a": A}


def load_module():
    spec = importlib.util.spec_from_file_location("translation_docx_helpers", HELPER)
    assert spec is not None and spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def qn(ns: str, name: str) -> str:
    return f"{{{ns}}}{name}"


def test_has_chinese() -> None:
    module = load_module()
    assert module.has_chinese("中文 ABC")
    assert not module.has_chinese("English only")


def test_insert_translation_after_body_paragraph_with_pair_marker_and_font() -> None:
    module = load_module()
    xml = f"""<w:document xmlns:w="{W}"><w:body>
      <w:p><w:pPr><w:jc w:val="both"/></w:pPr><w:r><w:t>这是中文段落。</w:t></w:r></w:p>
    </w:body></w:document>"""
    root = etree.fromstring(xml.encode("utf-8"))
    source = root.find(".//w:p", NS)
    inserted = module.insert_translation_paragraph_after(
        source,
        "This is an English paragraph.",
        pair_id="p0001",
        bookmark_start_id=10,
    )
    paras = root.findall(".//w:body/w:p", NS)
    assert paras == [source, inserted]
    assert "".join(inserted.xpath(".//w:t/text()", namespaces=NS)) == "This is an English paragraph."
    assert inserted.find("./w:pPr/w:jc", NS).get(qn(W, "val")) == "both"
    names = {item.get(qn(W, "name")) for item in root.findall(".//w:bookmarkStart", NS)}
    assert {"btx_p0001_src", "btx_p0001_en"} <= names
    run_fonts = inserted.find(".//w:rFonts", NS)
    assert run_fonts.get(qn(W, "ascii")) == "Times New Roman"
    assert run_fonts.get(qn(W, "hAnsi")) == "Times New Roman"
    assert run_fonts.get(qn(W, "cs")) == "Times New Roman"


def test_insert_translation_inside_same_table_cell() -> None:
    module = load_module()
    xml = f"""<w:document xmlns:w="{W}"><w:body><w:tbl><w:tr><w:tc>
      <w:p><w:r><w:t>表格中文</w:t></w:r></w:p>
    </w:tc></w:tr></w:tbl></w:body></w:document>"""
    root = etree.fromstring(xml.encode("utf-8"))
    source = root.find(".//w:tc/w:p", NS)
    inserted = module.insert_translation_paragraph_after(
        source,
        "English in the same cell",
        pair_id="t0001",
        bookmark_start_id=20,
    )
    assert root.findall(".//w:tc/w:p", NS) == [source, inserted]


def test_insert_translation_copies_omml_formula_into_english_pair() -> None:
    module = load_module()
    xml = f"""<w:document xmlns:w="{W}" xmlns:m="{M}"><w:body>
      <w:p><w:r><w:t>中文公式</w:t></w:r><m:oMath><m:r><m:t>x+1=2</m:t></m:r></m:oMath></w:p>
    </w:body></w:document>"""
    root = etree.fromstring(xml.encode("utf-8"))
    source = root.find(".//w:p", NS)
    inserted = module.insert_translation_paragraph_after(source, "English formula", "p0002")
    assert inserted.find(".//m:oMath", NS) is not None
    assert module.collect_formula_signatures(source) == module.collect_formula_signatures(inserted)


def test_insert_translation_copies_drawing_formula_and_rewrites_relationships() -> None:
    module = load_module()
    xml = f"""<w:document xmlns:w="{W}" xmlns:r="{R}" xmlns:a="{A}"><w:body>
      <w:p><w:r><w:t>图像公式</w:t></w:r><w:r><w:drawing><a:blip r:embed="rId5"/></w:drawing></w:r></w:p>
    </w:body></w:document>"""
    root = etree.fromstring(xml.encode("utf-8"))
    source = root.find(".//w:p", NS)
    inserted = module.insert_translation_paragraph_after(
        source,
        "Image formula",
        "p0003",
        relationship_id_map={"rId5": "rId42"},
    )
    assert inserted.xpath(".//@r:embed", namespaces=NS) == ["rId42"]
    assert inserted.find("./w:r/w:drawing", NS) is not None
    assert inserted.find("./w:drawing", NS) is None
    signatures = module.collect_formula_signatures(
        inserted,
        relationship_targets={"rId42": "media/image1.png"},
        related_blobs={"media/image1.png": b"formula-bytes"},
    )
    assert signatures[0]["type"] == "w:drawing"
    assert signatures[0]["target"] == "media/image1.png"
    assert "sha256" in signatures[0]


def test_insert_translation_copies_ole_and_vml_formula_runs_without_flattening() -> None:
    module = load_module()
    xml = f"""<w:document xmlns:w="{W}" xmlns:r="{R}" xmlns:v="{V}"><w:body>
      <w:p>
        <w:r><w:t>对象公式</w:t></w:r>
        <w:r><w:object r:id="rId7"/></w:r>
        <w:r><w:pict><v:shape><v:imagedata r:id="rId8"/></v:shape></w:pict></w:r>
      </w:p>
    </w:body></w:document>"""
    root = etree.fromstring(xml.encode("utf-8"))
    source = root.find(".//w:p", NS)
    inserted = module.insert_translation_paragraph_after(
        source,
        "Object formula",
        "p0004",
        relationship_id_map={"rId7": "rId70", "rId8": "rId80"},
    )
    assert inserted.find("./w:r/w:object", NS).get(qn(R, "id")) == "rId70"
    imagedata = inserted.find("./w:r/w:pict/v:shape/v:imagedata", NS)
    assert imagedata is not None
    assert imagedata.get(qn(R, "id")) == "rId80"
    assert inserted.find("./w:object", NS) is None
    assert inserted.find("./v:shape", NS) is None


def test_insert_translation_preserves_note_refs_and_markers_without_superscripting_body_text() -> None:
    module = load_module()
    xml = f"""<w:document xmlns:w="{W}"><w:body>
      <w:p>
        <w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:t>1</w:t></w:r>
        <w:r><w:rPr><w:b/></w:rPr><w:t>中文正文</w:t></w:r>
        <w:r><w:footnoteReference w:id="2"/></w:r>
        <w:r><w:endnoteReference w:id="3"/></w:r>
        <w:r><w:rPr><w:vertAlign w:val="subscript"/></w:rPr><w:t>i</w:t></w:r>
      </w:p>
    </w:body></w:document>"""
    root = etree.fromstring(xml.encode("utf-8"))
    source = root.find(".//w:p", NS)
    inserted = module.insert_translation_paragraph_after(source, "English body text", "p0004b")
    english_run = inserted.find("./w:r", NS)
    assert english_run is not None
    assert english_run.find("./w:rPr/w:b", NS) is not None
    assert english_run.find("./w:rPr/w:vertAlign", NS) is None
    assert inserted.xpath(".//w:footnoteReference/@w:id", namespaces=NS) == ["2"]
    assert inserted.xpath(".//w:endnoteReference/@w:id", namespaces=NS) == ["3"]
    assert inserted.xpath(".//w:vertAlign/@w:val", namespaces=NS) == ["superscript", "subscript"]


def test_generated_docx_opens_after_translation_insertion() -> None:
    module = load_module()
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        input_docx = tmp_path / "input.docx"
        output_docx = tmp_path / "output.docx"
        package_dir = tmp_path / "package"
        document = Document()
        document.add_paragraph("这是中文段落。")
        document.save(input_docx)

        with zipfile.ZipFile(input_docx) as archive:
            archive.extractall(package_dir)
        document_xml = package_dir / "word" / "document.xml"
        tree = etree.parse(str(document_xml))
        source = tree.getroot().find(".//w:body/w:p", NS)
        assert source is not None
        module.insert_translation_paragraph_after(source, "This is an English paragraph.", "p0005")
        tree.write(str(document_xml), encoding="UTF-8", xml_declaration=True)

        with zipfile.ZipFile(output_docx, "w", zipfile.ZIP_DEFLATED) as archive:
            for path in package_dir.rglob("*"):
                if path.is_file():
                    archive.write(path, path.relative_to(package_dir).as_posix())
        reopened = Document(output_docx)
        assert [paragraph.text for paragraph in reopened.paragraphs[:2]] == [
            "这是中文段落。",
            "This is an English paragraph.",
        ]


def test_copy_relationship_target_updates_rels_content_types_for_image_and_ole() -> None:
    module = load_module()
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        src = root / "src"
        dst = root / "dst"
        (src / "word" / "_rels").mkdir(parents=True)
        (src / "word" / "media").mkdir(parents=True)
        (src / "word" / "embeddings").mkdir(parents=True)
        (dst / "word" / "_rels").mkdir(parents=True)
        (src / "word" / "media" / "image1.png").write_bytes(b"formula-image")
        (src / "word" / "embeddings" / "oleObject1.bin").write_bytes(b"ole-formula")
        (src / "word" / "_rels" / "document.xml.rels").write_text(
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId5" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image1.png"/>'
            '<Relationship Id="rId7" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject" Target="embeddings/oleObject1.bin"/>'
            "</Relationships>",
            encoding="utf-8",
        )
        (dst / "word" / "_rels" / "document.xml.rels").write_text(
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>',
            encoding="utf-8",
        )
        content_types = '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"></Types>'
        (src / "[Content_Types].xml").write_text(content_types, encoding="utf-8")
        (dst / "[Content_Types].xml").write_text(content_types, encoding="utf-8")
        image_rid, image_target = module.copy_relationship_target(src, dst, "word/document.xml", "rId5")
        ole_rid, ole_target = module.copy_relationship_target(src, dst, "word/document.xml", "rId7")
        assert image_rid == "rId1"
        assert image_target == "media/image1.png"
        assert ole_rid == "rId2"
        assert ole_target == "embeddings/oleObject1.bin"
        assert (dst / "word" / "media" / "image1.png").read_bytes() == b"formula-image"
        assert (dst / "word" / "embeddings" / "oleObject1.bin").read_bytes() == b"ole-formula"
        rels_text = (dst / "word" / "_rels" / "document.xml.rels").read_text(encoding="utf-8")
        assert 'Id="rId1"' in rels_text
        assert 'Target="media/image1.png"' in rels_text
        assert 'Id="rId2"' in rels_text
        assert 'Target="embeddings/oleObject1.bin"' in rels_text
        content_type_text = (dst / "[Content_Types].xml").read_text(encoding="utf-8")
        assert 'Extension="png"' in content_type_text
        assert 'Extension="bin"' in content_type_text
        assert 'ContentType="application/vnd.openxmlformats-officedocument.oleObject"' in content_type_text


def test_copy_relationship_target_uses_unique_target_when_destination_exists() -> None:
    module = load_module()
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        src = root / "src"
        dst = root / "dst"
        (src / "word" / "_rels").mkdir(parents=True)
        (src / "word" / "media").mkdir(parents=True)
        (dst / "word" / "_rels").mkdir(parents=True)
        (dst / "word" / "media").mkdir(parents=True)
        (src / "word" / "media" / "image1.png").write_bytes(b"source-formula-image")
        (dst / "word" / "media" / "image1.png").write_bytes(b"existing-destination-image")
        (src / "word" / "_rels" / "document.xml.rels").write_text(
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId5" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/image1.png"/>'
            "</Relationships>",
            encoding="utf-8",
        )
        (dst / "word" / "_rels" / "document.xml.rels").write_text(
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>',
            encoding="utf-8",
        )
        content_types = '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"></Types>'
        (src / "[Content_Types].xml").write_text(content_types, encoding="utf-8")
        (dst / "[Content_Types].xml").write_text(content_types, encoding="utf-8")

        new_rid, new_target = module.copy_relationship_target(src, dst, "word/document.xml", "rId5")

        assert new_rid == "rId1"
        assert new_target == "media/image2.png"
        assert (dst / "word" / "media" / "image1.png").read_bytes() == b"existing-destination-image"
        assert (dst / "word" / "media" / "image2.png").read_bytes() == b"source-formula-image"
        rels_text = (dst / "word" / "_rels" / "document.xml.rels").read_text(encoding="utf-8")
        assert 'Target="media/image2.png"' in rels_text
