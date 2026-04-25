#!/usr/bin/env python3
from __future__ import annotations

import importlib.util
import sys
import tempfile
import zipfile
from pathlib import Path

try:
    from docx import Document
    from lxml import etree
except ImportError:
    print("Missing dependency: install python-docx and lxml")
    sys.exit(1)


ROOT = Path(__file__).resolve().parents[2]
HELPER = ROOT / "skills" / "word-default-formatting" / "formatting_helpers.py"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}


def load_module():
    spec = importlib.util.spec_from_file_location("formatting_helpers", HELPER)
    if spec is None or spec.loader is None:
        raise RuntimeError("cannot load formatting_helpers.py")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def attr(name: str) -> str:
    return f"{{{W}}}{name}"


def main() -> None:
    if not HELPER.exists():
        print(f"Missing helper: {HELPER}")
        sys.exit(1)

    with tempfile.TemporaryDirectory() as tmp:
        output_docx = Path(tmp) / "formatted.docx"
        module = load_module()

        doc = Document()
        module.set_style_mixed_fonts(doc.styles["Normal"])
        run = doc.add_paragraph().add_run("中文 ABC 123")
        module.set_run_mixed_fonts(run)
        doc.save(output_docx)

        with zipfile.ZipFile(output_docx) as archive:
            document_xml = archive.read("word/document.xml")
            styles_xml = archive.read("word/styles.xml")

        document_root = etree.fromstring(document_xml)
        run_fonts = document_root.find(".//w:rFonts", NS)
        if run_fonts is None:
            print("Run is missing w:rFonts")
            sys.exit(1)
        if run_fonts.get(attr("ascii")) != "Times New Roman":
            print("Run is missing Times New Roman ascii font")
            sys.exit(1)
        if run_fonts.get(attr("hAnsi")) != "Times New Roman":
            print("Run is missing Times New Roman hAnsi font")
            sys.exit(1)
        if run_fonts.get(attr("eastAsia")) != "宋体":
            print("Run is missing Songti eastAsia font")
            sys.exit(1)

        styles_root = etree.fromstring(styles_xml)
        normal_style = styles_root.find(".//w:style[@w:styleId='Normal']", NS)
        style_fonts = normal_style.find(".//w:rFonts", NS) if normal_style is not None else None
        if style_fonts is None:
            print("Normal style is missing w:rFonts")
            sys.exit(1)
        if style_fonts.get(attr("ascii")) != "Times New Roman":
            print("Normal style is missing Times New Roman ascii font")
            sys.exit(1)
        if style_fonts.get(attr("eastAsia")) != "宋体":
            print("Normal style is missing Songti eastAsia font")
            sys.exit(1)

    print("PASS: default formatting smoke test")


def test_default_formatting_smoke() -> None:
    main()


if __name__ == "__main__":
    main()
